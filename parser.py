#!/usr/bin/env python
#
# Copyright (c) 2018 - present.  Boling Consulting Solutions (bcsw.net)
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
# http://www.apache.org/licenses/LICENSE-2.0
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.
#
from __future__ import (
    absolute_import, division, print_function, unicode_literals
)
import argparse
import time
import copy
from docx import Document

from parser_lib.class_id import ClassIdList
from parser_lib.parsed_json import ParsedJson
from parser_lib.parsed_yaml import MetadataYAML
from preparsed_json import PreParsedJson
from parser_lib.versions import VersionHeading
from parser_lib.text import camelcase

#
#  This application takes the pre-parsed JSON file and further parses it to output suitable for
#  running a code generator with.
#
#  The process from scratch is:
#
#   1. Pre-parse the G.988 Word Document via 'preParse.py' to create the 'G.988.Precompiled.json' file
#
#   2. Parse the G.988 JSON via 'parser.py' to create the G.988.Parsed.json file.  At this point,
#      there is just a minimal fragment 'G.988.augment.yaml' file that really as a little bit of data
#
#   3. Run the 'augmentGenerator.py' file to create an 'augmented.yaml' file in the 'metadata' sub-
#      directory. This will have all the newly parsed JSON converted to YAML form.
#
#   4. Hand edit the augmented.yaml file by hand and adjust the values as needed.  Once you are done,
#      overwrite the base directory's 'G.988.augment.yaml' file with this.  You now have a
#      metadata file that will be used every time you either run the parser (or code generator) with
#      your updated metadata.
#
#   5. Run the 'parser.py' program again. This will pick up the metadata hint you created in step 4
#      and will create an updated 'G.988.Parsed.json' file with your data.
#
#   6. Run either the C or Go code generator to generate your classes.
#
############################################################################
#
#   If the pre-parser is upgraded or a bug is fixed and needs to be ran again due to updates, run steps 1-6
#
#   If the parser is upgraded or a bug is fixed, you only need to run steps 5 & 6
#
#   If the code generator you are interested in is upgraded or fixed, you only need to run step 6
#
############################################################################

MEClassSection = "11.2.4"       # Class IDs


def parse_args():
    parser = argparse.ArgumentParser(description='G.988 Parser')

    parser.add_argument('--ITU', '-I', action='store',
                        default='T-REC-G.988-201711-I!!MSW-E.docx',
                        help='Path to ITU G.988 specification document')

    parser.add_argument('--input', '-i', action='store',
                        default='G.988.PreCompiled.json',
                        help='Path to pre-parsed G.988 data, default: G.988.PreCompiled.json')

    parser.add_argument('--hints', '-H', action='store',
                        default='G.988.augment.yaml',
                        help='Path to hand modified G.988 input data, default: G.988.augment.yaml')

    parser.add_argument('--output', '-o', action='store',
                        default='G.988.Parsed.json',
                        help='Output filename, default: G.988.Parsed.json')

    parser.add_argument('--classes', '-c', action='store',
                        default='11.2.4',
                        help='Document section number with ME Class IDs, default: 11.2.4')

    args = parser.parse_args()
    return args


class Main(object):
    """ Main program """
    def __init__(self):
        self.args = parse_args()
        self.paragraphs = None
        self.body = None

        self.preparsed = PreParsedJson()
        self.parsed = ParsedJson()
        self.attribute_hints = dict()       # Class ID -> attribute hints list
        version = VersionHeading()
        version.name = 'parser'
        version.create_time = time.time()
        version.itu_document = self.args.ITU
        version.version = self.get_version()
        version.sha256 = self.get_file_hash(version.itu_document)
        self.parsed.add(version)

    @staticmethod
    def get_version():
        with open('VERSION', 'r') as f:
            for line in f:
                line = line.strip().lower()
                if len(line) > 0:
                    return line

    @staticmethod
    def get_file_hash(filename):
        import hashlib
        with open(filename, 'rb') as f:
            data = f.read()
            return hashlib.sha256(data).hexdigest()

    @property
    def sections(self):
        return self.preparsed.sections

    @property
    def section_list(self):
        return self.preparsed.section_list

    @property
    def class_ids(self):
        return self.parsed.class_ids

    def load_itu_document(self):
        return Document(self.args.ITU)

    def start(self):
        print("Loading ITU Document '{}' and parsed data file '{}'".format(self.args.ITU,
                                                                           self.args.input))
        self.preparsed.load(self.args.input)
        for version in self.preparsed.versions:
            self.parsed.add(version)

        self.attribute_hints = MetadataYAML().load(self.args.hints)

        document = self.load_itu_document()
        self.paragraphs = document.paragraphs
        # doc_sections = document.sections
        # styles = document.styles
        # self.body = document.element.body

        print('Extracting ME Class ID values')
        class_ids = ClassIdList.parse_sections(self.section_list,
                                               self.args.classes)
        for _cid, me_class in class_ids.items():
            self.parsed.add(me_class)

        print('Found {} ME Class ID entries. {} have sections associated to them'.
              format(len(self.class_ids),
                     len([c for c in self.class_ids.values()
                          if c.section is not None])))

        # TODO: These need more work. skipping for now
        crazy_formatted_mes = \
            {23, 319, 320,  # CES physical interface performance MEs
             164,           # MoCA interface performance
             165,           # VDLS2 line config extensions
             157,           # Large String                      (part of AT&T OpenOMCI v3.0)
             415}

        print('Skipping the following MEs due to complex document formatting')
        print("    {}".format(crazy_formatted_mes))
        todo_class_ids = {k: v for k, v in self.class_ids.items()
                          if k not in crazy_formatted_mes}

        print('Managed Entities without Sections')
        for c in [c for c in todo_class_ids.values() if c.section is None]:
            print('    {:>4}: {}'.format(c.cid, c.name))

        # Work with what we have
        todo_class_ids = {cid: c for cid, c in todo_class_ids.items()
                          if c.section is not None}
        print('')
        print('Parsing deeper for managed Entities with Sections')

        final_class_ids = todo_class_ids

        for c in final_class_ids.values():
            if c.section is None:
                c.failure(None, None)
                continue

            print('    {:>9}:  {:>4}: {} -> {}'.format(c.section.section_number,
                                                       c.cid,
                                                       c.name,
                                                       camelcase(c.name)))
            c.deep_parse(self.paragraphs)
            pass

        # Some just need some manual intervention
        final_class_ids = self.fix_difficult_class_ids(final_class_ids)

        # Who creates them
        final_class_ids = self.find_class_access(final_class_ids)

        completed = len([c for c in final_class_ids.values() if c.state == 'complete'])
        failed = len([c for c in final_class_ids.values() if c.state == 'failure'])

        print('Of {} MEs, {} were parsed successfully and {} failed'.format(len(final_class_ids),
                                                                            completed,
                                                                            failed))
        # Run some sanity checks
        print('\n\n\nValidating ME Class Information, total of {}:\n'.
              format(len(final_class_ids)))

        class_with_issues = dict()
        class_with_no_actions = dict()
        class_with_no_attributes = dict()
        attributes_with_no_access = dict()
        attributes_with_no_size = dict()
        attributes_with_zero_size = dict()
        class_with_too_many_attributes = dict()
        num_attributes = 0

        for c in final_class_ids.values():
            print('  ID: {}: {} -\t{}'.format(c.cid, c.section.section_number, c.name),
                  end='')

            if c.state != 'complete':
                print('\t\tParsing ended in state {}', c.state)
                class_with_issues[c.cid] = c

            if len(c.actions) == 0:
                print('\t\tActions: No actions decoded for ME')
                class_with_issues[c.cid] = c
                class_with_no_actions[c.cid] = c
                c.failure(None, None)       # Mark invalid
            else:
                print('\t\tActions: {}'.format({a.name for a in c.actions}))

            if len(c.attributes) == 0:
                print('\t\tNO ATTRIBUTES')      # TODO Look for 'set' without 'get'
                class_with_issues[c.cid] = c
                class_with_no_attributes[c.cid] = c
                c.failure(None, None)       # Mark invalid

            elif len(c.attributes) > 17:        # Entity ID counts as well in this list
                print('\t\tTOO MANY ATTRIBUTES')
                class_with_issues[c.cid] = c
                class_with_too_many_attributes[c.cid] = c
                c.failure(None, None)       # Mark invalid

            else:
                for attr in c.attributes:
                    num_attributes += 1
                    print('\t\t\t\t{}'.format(attr.name), end='')
                    if attr.access is None or len(attr.access) == 0:
                        print('\t\t\t\tNO ACCESS INFORMATION')
                        attributes_with_no_access[c.cid] = c
                        c.failure(None, None)       # Mark invalid
                    else:
                        print('\t\t\t\tAccess: {}'.format({a.name for a in attr.access}))

                    if attr.size is None:
                        attributes_with_no_size[c.cid] = c
                        print('\t\t\t\tNO SIZE INFORMATION')
                        c.failure(None, None)       # Mark invalid

                    elif attr.size.octets == 0:
                        attributes_with_zero_size[c.cid] = c
                        print('\t\t\t\tSIZE zero')
                        c.failure(None, None)       # Mark invalid

        print('Section parsing is complete, saving JSON output...')
        print('=======================================================')

        # Output the results to a JSON file so it can be used by a code-generation
        # tool
        self.parsed.save(self.args.output)

        # Restore and verify
        self.parsed.load(self.args.output)
        self.parsed.dump()

        # Results
        print("Of the {} class IDs, {} had issues: {} had no actions and {} had no attributes and {} with too many".
              format(len(final_class_ids), len(class_with_issues), len(class_with_no_actions),
                     len(class_with_no_attributes), len(class_with_too_many_attributes)))

        print("Of the {} attributes, {} had no access info and {} had no size info and {} had zero size".
              format(num_attributes, len(attributes_with_no_access), len(attributes_with_no_size),
                     len(attributes_with_zero_size)))

    def fix_difficult_class_ids(self, class_list):
        # Special exception. Ethernet frame performance monitoring history data downstream
        # is in identical upstream and only a note of that exists. Fix it now
        from parser_lib.actions import Actions
        from parser_lib.size import AttributeSize
        from parser_lib.attributes import AttributeAccess

        # For SIP user data, the Username&Password attribute is a pointer
        # to a security methods ME and is 2 bytes but is in the document as
        # just (2)
        if 153 in class_list.keys():
            sip = class_list[153]
            sz = AttributeSize()
            sz._octets = 2
            sip.attributes[4].size = sz

        # ONU remote debug - reply table size not bounded
        if 158 in class_list.keys():
            item = class_list[158]
            reply_table = item.attributes[3]
            sz = copy.deepcopy(reply_table.size)
            sz._octets = -1
            reply_table.size = sz

        # MCAST GEM Interworking - IPv4
        if 281 in class_list.keys():
            item = class_list[281]
            sz = AttributeSize()
            sz._octets = 12
            sz.getnext_required = True
            item.attributes[9].size = sz
            item.attributes[9].access.add(AttributeAccess.Read)
            item.attributes[9].access.add(AttributeAccess.Write)

        # OMCI.  IPv6 Table is 24N octets
        if 287 in class_list.keys():
            item = class_list[287]
            item.attributes[1].getnext_required = True

            sz = AttributeSize()
            sz._octets = 1
            sz.getnext_required = True
            item.attributes[2].size = sz

        # Managed entity tables.  4 tables need fixing
        if 288 in class_list.keys():
            me = class_list[288]
            class_list[288].name += ' ME'       # To avoid conflicts with Go file/struct names
            sz = AttributeSize()
            sz._octets = 1
            me.attributes[4].size = sz
            me.attributes[5].size = sz

        # Managed entity code points table.  Table is 2*n octets
        if 289 in class_list.keys():
            class_list[289].name += ' ME'       # To avoid conflicts with Go file/struct names

        # Dot1ag maintenance domain
        if 299 in class_list.keys():
            item = class_list[299]
            item.attributes[3].size._repeat_count = 2
            item.attributes[3].size._repeat_max = 2

        # Dot1ag maintenance domain
        if 300 in class_list.keys():
            item = class_list[300]
            item.attributes[3].size._repeat_count = 2
            item.attributes[3].size._repeat_max = 2
            sz = AttributeSize()
            sz._octets = 24
            item.attributes[5].size = sz

        # Octet Stirng - 1..15 row entries
        if 307 in class_list.keys():
            item = class_list[307]
            part1 = item.attributes[2]
            part1.name = 'Part 1 to 15'
            part1.optional = False
            part1.size._repeat_max = 15

        # General Purpose Buffer - Buffer table is one big unbounded string
        if 308 in class_list.keys():
            item = class_list[308]
            buffer_table = item.attributes[2]
            sz = copy.deepcopy(buffer_table.size)
            sz._octets = -1
            buffer_table.size = sz

        # For multicast operations profile - Attributes getting polluted with table info/descriptions
        if 309 in class_list.keys():
            item = class_list[309]
            if len(item.attributes) == 20:
                first_bad = 8   # Table Control is description of attribute 7
                next_good = 11  # Pick back up to good attributes at the Static Access Coltrol List Table
                # Also attribute 7 lost its information on access...
                
                item.attributes = item.attributes[0:first_bad] + item.attributes[next_good:]
                # And a quick attribute name fixup here
                item.attributes[16].name = "Downstream IGMP and multicast TCI"

        # For multicast subscriber config info. very hard to decode automatically
        if 310 in class_list.keys():
            msci = class_list[310]
            msci.attributes.remove(8)
            item = msci.attributes[7]
            sz = AttributeSize()
            sz._octets = 22
            sz.getnext_required = True
            item.size = sz
            item.access.add(AttributeAccess.Read)
            item.access.add(AttributeAccess.Write)

        if 321 in class_list.keys() and 322 in class_list.keys():
            down = class_list[321]
            up = class_list[322]
            down.attributes = up.attributes
            down.actions = up.actions
            down.optional_actions = up.optional_actions
            down.alarms = up.alarms
            down.avcs = up.avcs
            down.test_results = up.test_results
            down.hidden = up.hidden

        # xDSL line inventory and status data part 5
        if 325 in class_list.keys():
            item = class_list[325]
            try:
                from parser_lib.actions import Actions
                # Type in document, not table attributes present
                item.actions.remove(Actions.GetNext)
            except KeyError:
                pass

        # Enhanced security control
        if 332 in class_list.keys():
            item = class_list[332]
            try:
                # Type in document, not table attributes present
                item.attributes[7].access.add(AttributeAccess.Read)
            except KeyError:
                pass

        # xDSL line inventory and status data part 8
        if 414 in class_list.keys():
            item = class_list[414]
            try:
                # Type in document, not table attributes present
                item.actions.remove(Actions.GetNext)
            except KeyError:
                pass

        # Ethernet frame extended PM 64-bit. fix actions
        if 426 in class_list.keys():
            item = class_list[426]
            item.actions.update([Actions.Create, Actions.Delete,
                                Actions.Get, Actions.Set])
            item.optional_actions.add(Actions.GetCurrentData)
            try:
                if 334 in class_list.keys():
                    extended32 = class_list[334]
                    extended64 = class_list[426]
                    extended64.attributes = copy.deepcopy(extended32.attributes)
                    extended64.actions = extended32.actions
                    extended64.optional_actions = extended32.optional_actions
                    extended64.alarms = extended32.alarms
                    extended64.avcs = extended32.avcs
                    extended64.test_results = extended32.test_results
                    extended64.hidden = extended32.hidden

                    sz = AttributeSize()
                    sz._octets = 8
                    for index in range(3, 17):
                        extended64.attributes[index].size = sz

                pass    # TODO: Also no attributes are getting decoded
            except KeyError:
                pass

        # Now even some other crazy things
        class_list = self.fix_other_difficulties(class_list)

        # Find counter and other types of attributes
        self.find_attribute_types(class_list)

        return class_list

    @staticmethod
    def fix_other_difficulties(class_list):
        # Some uncommon cleanups
        # for cid, cls in class_list.items():
        #     pass
        return class_list

    @staticmethod
    def find_class_access(class_list):
        from parser_lib.class_id import ClassAccess, Actions
        for _, item in class_list.items():
            if item.access == ClassAccess.UnknownAccess and len(item.actions):
                if Actions.Create in item.actions:
                    item.access = ClassAccess.CreatedByOlt
                else:
                    item.access = ClassAccess.CreatedByOnu

        return class_list

    def find_attribute_types(self, class_list):
        # A bit more in depth look at the attributes
        from parser_lib.attributes import AttributeType

        for cid, item in class_list.items():
            # Any hints to apply before seeding default settings?
            if cid in self.attribute_hints:
                hints = self.attribute_hints[cid]
                for index, attr_hint in hints['attributes'].items():
                    item_attribute = next((attr for attr in item.attributes if attr.index == index), None)
                    if item_attribute is not None:
                        if 'type' in attr_hint:
                            item_attribute.attribute_type = attr_hint['type']
                        if 'default' in attr_hint:
                            item_attribute.default = attr_hint['default']
                        if 'constraint' in attr_hint:
                            item_attribute.constraint = attr_hint['constraint']

            for attr in item.attributes:
                # Only set unknown types (or integer since it may be a counter)
                if attr.attribute_type == AttributeType.Unknown:
                    attr.find_type(item)

        return class_list


if __name__ == '__main__':
    try:
        Main().start()

    except Exception as _e:
        raise
